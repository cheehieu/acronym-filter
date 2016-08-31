#!/usr/bin/python

'''
acronym-filter.py
Written by Hieu Nguyen (hnguyen@intuitivemachines.com) on 8/31/16

This acronym-filter script was written to automate the process of generating a glossary of acronyms used during the Axiom Module system requirements review (SRR) on September 1, 2016. It scans a Word .DOCX document and puts every unique acronym (a consecutive sequence of capital letters, symbols, or numbers) into a dictionary. If applicable, the acronyms are matched to an imported .CSV dictionary file to retrieve the acronym expansions. The newly found list of acronyms is saved to an output_dict.csv dictionary file, where it can be manually modified with Excel. An output_glossary.docx file is also generated from the input dictionary, which can then be used for final publishing of an acronym glossary. 

Note: you may need to install lxml and python-docx dependencies.
	$ pip install lxml
	$ pip install python-docx

Limitations:
* Acronyms must begin with a capital letter and end with another captial letter or number. Thus, lowercase abbreviations such as units (mm, dBA, etc.) will not be captured by the script. But ACRONYM_REGX can be modified to change this search pattern.
* Long acronyms are not captured. A workaround was used because the IRD requirements were written using all caps. Change MAX_ACRO_LENGTH to tune.
* Identical acronyms (with different expansions) are not captured. The script filters duplicate acronyms without using any contextual differentiation.

Useful Resources:
* http://spaceflight.nasa.gov/cgi-bin/acronyms.cgi
* http://www.nasa.gov/directorates/heo/scan/definitions/acronyms/index.html
'''


import sys
sys.path.append('/usr/local/lib/python2.7/site-packages')
from docx import Document
import re
import csv
from collections import OrderedDict


# Defines and flags
ACRONYM_REGX = '([A-Z][a-zA-Z0-9+\.\&\-]*[A-Z0-9])'
MAX_ACRO_LENGTH = 7	# used to filter out false acronyms from IRD reqs
GENERATE_GLOSSARY_FROM_DICTIONARY = True	# set True to generate .DOCX

# Get command line input arguments
if len(sys.argv) < 2:
	print "\n\tUSAGE: "
	print "\t$ python acronym-filter INPUT_FILE.docx INPUT_DICT.csv\n\n"
	sys.exit(0)
else:
	DOCX_FILENAME = sys.argv[1]
	print "\nInput document:\t\t" + DOCX_FILENAME
	IMPORT_DICTIONARY = False
	if len(sys.argv) > 2:
		DICT_FILENAME = sys.argv[2]
		print "Input dictionary:\t" + DICT_FILENAME
		IMPORT_DICTIONARY = True
	print ""

# Search document (paragraphs and table cells) for acronyms
full_acro_list = []
document = Document(DOCX_FILENAME)
for paragraph in document.paragraphs:
    searchObj = re.findall(ACRONYM_REGX, paragraph.text)
    for ii in xrange(len(searchObj)):
		full_acro_list.append(searchObj[ii])
for table in document.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                searchObj = re.findall(ACRONYM_REGX, paragraph.text)
            	for ii in xrange(len(searchObj)):
		    		full_acro_list.append(searchObj[ii])

# Clean up acronym list by removing duplicates
unique_acro_list = list(set(full_acro_list))

# Attempt to remove false acronyms from IRD req titles
valid_acro_list = [s for s in unique_acro_list if len(s) <= MAX_ACRO_LENGTH]

# Add acronyms to a dictionary
acro_dict = dict((key,"") for key in valid_acro_list)

# Import values from existing dictionary in CSV file
if (IMPORT_DICTIONARY):
	with open(DICT_FILENAME, 'rbU') as csv_file:
		reader = csv.reader(csv_file, delimiter=',', dialect=csv.excel_tab)
		import_dict = dict(reader)

	# Pull in acronym expanded form values from imported dictionary
	for key, value in import_dict.iteritems():
		if key in acro_dict:
			acro_dict[key] = value

# Sort dictionary alphabetically
sorted_acro_dict = OrderedDict(sorted(acro_dict.items(), key=lambda t: t[0]))

# Print acronyms and expanded values to console
for key in sorted_acro_dict:
	print key, " - ", sorted_acro_dict[key]

# Save acronyms and expanded values to a new dictionary .CSV file
with open('output_dict.csv', 'wb') as csv_file:
    writer = csv.writer(csv_file)
    for key, value in sorted_acro_dict.items():
       writer.writerow([key, value])

# Save old dictionary to a table in .DOCX file
if (IMPORT_DICTIONARY and GENERATE_GLOSSARY_FROM_DICTIONARY):
	sorted_import_dict = OrderedDict(sorted(import_dict.items(), key=lambda t: t[0]))
	document = Document()
	table = document.add_table(rows=1, cols=2)
	hdr_cells = table.rows[0].cells
	hdr_cells[0].text = 'Abbreviation'
	hdr_cells[1].text = 'Expansion'
	for key in sorted_import_dict:
		row_cells = table.add_row().cells
		row_cells[0].text = key
		row_cells[1].text = sorted_import_dict[key]
	document.save('output_glossary.docx')

# Print acronym-filter stats
print "\n\nTotal # of Acronyms:\t\t", len(full_acro_list)
print "Total # of Unique Acronyms:\t", len(unique_acro_list)
print "Total # of Valid Acronyms:\t", len(valid_acro_list)
if (IMPORT_DICTIONARY):
	print "Imported Dictionary Length:\t", len(import_dict)
	if (GENERATE_GLOSSARY_FROM_DICTIONARY):
		print "A new glossary was generated from %s and written to output_glossary.docx" % DICT_FILENAME
print "A new dictionary of Valid Acronyms was written to output_dict.csv"
print "\n"
